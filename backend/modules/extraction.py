"""
Multi-format text extraction.
Returns list of (section_title | None, text) tuples.

Image types (png/jpg/webp/gif/bmp) return empty list — 
caller (ingestion.py) routes them directly to OCR.
"""
import csv
import io
import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

IMAGE_TYPES = {"png", "jpg", "jpeg", "webp", "gif", "bmp", "tiff", "tif"}


def is_image_type(file_type: str) -> bool:
    return file_type.lower().lstrip(".") in IMAGE_TYPES


def extract(file_path: str, file_type: str) -> list[tuple[str | None, str]]:
    """Dispatch to format-specific extractor."""
    ext = file_type.lower().lstrip(".")
    p = Path(file_path)

    if ext in IMAGE_TYPES:
        # Images handled separately via OCR in ingestion.py
        return []

    try:
        match ext:
            case "pdf":
                return _pdf(p)
            case "docx":
                return _docx(p)
            case "html" | "htm":
                return _html(p)
            case "txt" | "md":
                return _plain(p)
            case "csv":
                return _csv(p)
            case "json":
                return _json(p)
            case _:
                return _plain(p)
    except Exception as e:
        logger.error(f"Extraction failed [{ext}] {file_path}: {e}", exc_info=True)
        raise


def _pdf(path: Path) -> list[tuple[str | None, str]]:
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    results = []
    for i, page in enumerate(doc):
        blocks = page.get_text("blocks")
        texts = [b[4].strip() for b in blocks if b[6] == 0 and b[4].strip()]
        if texts:
            results.append((f"Page {i + 1}", "\n".join(texts)))
    doc.close()
    return results


def _docx(path: Path) -> list[tuple[str | None, str]]:
    from docx import Document

    doc = Document(str(path))
    sections: list[tuple[str | None, str]] = []
    heading: str | None = None
    buf: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if "heading" in para.style.name.lower():
            if buf:
                sections.append((heading, "\n".join(buf)))
                buf = []
            heading = text
        else:
            buf.append(text)

    if buf:
        sections.append((heading, "\n".join(buf)))

    for table in doc.tables:
        rows = [
            " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            for row in table.rows
        ]
        rows = [r for r in rows if r]
        if rows:
            sections.append(("Table", "\n".join(rows)))

    return sections or [(None, "")]


def _html(path: Path) -> list[tuple[str | None, str]]:
    from bs4 import BeautifulSoup
    import chardet

    raw = path.read_bytes()
    enc = chardet.detect(raw)["encoding"] or "utf-8"
    soup = BeautifulSoup(raw.decode(enc, errors="replace"), "lxml")

    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    sections: list[tuple[str | None, str]] = []
    heading: str | None = None
    buf: list[str] = []

    for el in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "pre", "code"]):
        text = el.get_text(separator=" ", strip=True)
        if not text:
            continue
        if el.name in ("h1", "h2", "h3", "h4"):
            if buf:
                sections.append((heading, "\n".join(buf)))
                buf = []
            heading = text
        else:
            buf.append(text)

    if buf:
        sections.append((heading, "\n".join(buf)))

    return sections or [(None, soup.get_text(separator="\n", strip=True))]


def _plain(path: Path) -> list[tuple[str | None, str]]:
    import chardet

    raw = path.read_bytes()
    enc = chardet.detect(raw)["encoding"] or "utf-8"
    return [(None, raw.decode(enc, errors="replace"))]


def _csv(path: Path) -> list[tuple[str | None, str]]:
    import chardet

    raw = path.read_bytes()
    enc = chardet.detect(raw)["encoding"] or "utf-8"
    reader = csv.reader(io.StringIO(raw.decode(enc, errors="replace")))
    rows = [" | ".join(row) for row in reader if any(c.strip() for c in row)]
    return [("CSV Data", "\n".join(rows))]


def _json(path: Path) -> list[tuple[str | None, str]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    try:
        pretty = json.dumps(json.loads(text), indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        pretty = text
    return [("JSON Data", pretty)]
